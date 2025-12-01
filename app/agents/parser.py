import asyncio
from dataclasses import dataclass
from typing import List, Optional

import docx
from pypdf import PdfReader


@dataclass
class DocumentContent:
    text: str
    metadata: dict
    tables: Optional[List[str]] = None  # tables rendered as markdown-ish strings


class DocumentParserAgent:
    """Deterministic parser for PDF/DOCX; no LLM involvement."""

    def parse_pdf(self, file_path: str) -> DocumentContent:
        """Parse PDF to text; best-effort per page."""
        reader = PdfReader(file_path)
        pages = []
        for idx, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            pages.append(f"[Page {idx}]\n{page_text.strip()}")
        text = "\n\n".join(pages)
        metadata = reader.metadata if reader.metadata else {}
        return DocumentContent(text=text, metadata=metadata, tables=None)

    def parse_docx(self, file_path: str) -> DocumentContent:
        """Parse DOCX paragraphs and tables."""
        doc = docx.Document(file_path)
        paras = []
        for para in doc.paragraphs:
            if para.text.strip():
                paras.append(para.text)
        text = "\n".join(paras)
        tables_md: List[str] = []
        for tbl in doc.tables:
            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                header, *body = rows
                if header:
                    md_header = "| " + " | ".join(header) + " |"
                    md_sep = "| " + " | ".join(["---"] * len(header)) + " |"
                    md_rows = ["| " + " | ".join(r) + " |" for r in body]
                    tables_md.append("\n".join([md_header, md_sep, *md_rows]))
        return DocumentContent(text=text, metadata={}, tables=tables_md or None)

    def parse_file(self, file_path: str) -> DocumentContent:
        if file_path.endswith(".pdf"):
            return self.parse_pdf(file_path)
        if file_path.endswith(".docx"):
            return self.parse_docx(file_path)
        raise ValueError("Unsupported file format")

    async def parse_file_async(self, file_path: str) -> DocumentContent:
        """Async wrapper to avoid blocking the event loop."""
        return await asyncio.to_thread(self.parse_file, file_path)
